#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Скрипт для конвертации ДОКЛАД.md в формат DOCX с сохранением форматирования
Улучшенная версия с поддержкой формул LaTeX через изображения
"""

import os
import sys
import re
import io
import tempfile
import shutil

# Настройка кодировки для Windows
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

def latex_to_image(latex_formula, is_block=False):
    """Конвертирует LaTeX формулу в изображение"""
    try:
        import matplotlib
        matplotlib.use('Agg')  # Non-interactive backend
        import matplotlib.pyplot as plt
        from matplotlib import mathtext
        
        # Настройка параметров для рендеринга
        fig = plt.figure(figsize=(10, 0.5) if is_block else (8, 0.3))
        fig.patch.set_facecolor('white')
        fig.patch.set_alpha(0)
        
        # Рендерим формулу
        if is_block:
            # Блочная формула - центрируем и делаем больше
            plt.text(0.5, 0.5, f'${latex_formula}$', 
                    fontsize=14, ha='center', va='center',
                    transform=fig.transFigure)
        else:
            # Inline формула
            plt.text(0.5, 0.5, f'${latex_formula}$', 
                    fontsize=12, ha='center', va='center',
                    transform=fig.transFigure)
        
        plt.axis('off')
        
        # Сохраняем во временный файл
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        temp_path = temp_file.name
        temp_file.close()
        
        plt.savefig(temp_path, dpi=150, bbox_inches='tight', 
                   pad_inches=0.1, facecolor='white', transparent=True)
        plt.close(fig)
        
        return temp_path
    except Exception as e:
        print(f"[WARNING] Не удалось создать изображение формулы: {e}")
        return None

def convert_markdown_to_docx():
    """Конвертирует ДОКЛАД.md в ДОКЛАД.docx"""
    
    input_file = 'ДОКЛАД.md'
    output_file = 'ДОКЛАД.docx'
    
    if not os.path.exists(input_file):
        print(f"Файл {input_file} не найден!")
        return False
    
    print(f"Конвертация {input_file} в {output_file}...")
    
    # Попробуем использовать pypandoc (требует установки pandoc) - лучший вариант
    try:
        import pypandoc
        print("Используется pypandoc (лучший вариант для формул)...")
        pypandoc.convert_file(
            input_file,
            'docx',
            outputfile=output_file,
            extra_args=[
                '--mathml',  # Используем MathML для формул (лучше для Word)
                '--standalone'
            ]
        )
        print(f"[OK] Файл {output_file} успешно создан с поддержкой формул!")
        return True
    except ImportError:
        print("pypandoc не установлен, пробуем альтернативный метод с изображениями...")
    except Exception as e:
        print(f"Ошибка при использовании pypandoc: {e}")
        print("Пробуем альтернативный метод с изображениями...")
    
    # Альтернативный метод: используем python-docx с конвертацией формул в изображения
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        print("Используется python-docx с конвертацией формул в изображения...")
        
        # Читаем markdown файл
        with open(input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # Создаем документ
        doc = Document()
        
        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        i = 0
        in_code_block = False
        in_math_block = False
        math_formula = ""
        temp_images = []  # Список временных файлов для удаления
        
        while i < len(lines):
            line = lines[i].rstrip('\n\r')
            original_line = line
            line_stripped = line.strip()
            
            # Пропускаем пустые строки (но добавляем параграф для разделения)
            if not line_stripped:
                if i > 0 and i < len(lines) - 1:
                    # Проверяем, не является ли это разделителем
                    if i + 1 < len(lines) and lines[i+1].strip().startswith('---'):
                        i += 1
                        continue
                doc.add_paragraph()
                i += 1
                continue
            
            # Блочные формулы LaTeX
            if line_stripped == '$$':
                if not in_math_block:
                    in_math_block = True
                    math_formula = ""
                    i += 1
                    continue
                else:
                    # Конец блочной формулы - создаем изображение
                    if math_formula:
                        img_path = latex_to_image(math_formula, is_block=True)
                        if img_path:
                            temp_images.append(img_path)
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(img_path, width=Inches(5))
                        else:
                            # Если не удалось создать изображение, добавляем как текст
                            p = doc.add_paragraph()
                            run = p.add_run(math_formula)
                            run.font.name = 'Cambria Math'
                            run.font.size = Pt(12)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    in_math_block = False
                    math_formula = ""
                    i += 1
                    continue
            
            if in_math_block:
                # Собираем формулу (может быть многострочной)
                if math_formula:
                    math_formula += " " + line_stripped
                else:
                    math_formula = line_stripped
                i += 1
                continue
            
            # Код блоки
            if line_stripped.startswith('```'):
                in_code_block = not in_code_block
                i += 1
                continue
            
            if in_code_block:
                p = doc.add_paragraph()
                run = p.add_run(line_stripped)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                i += 1
                continue
            
            # Заголовки
            if line_stripped.startswith('# '):
                p = doc.add_heading(line_stripped[2:], level=1)
                i += 1
                continue
            elif line_stripped.startswith('## '):
                p = doc.add_heading(line_stripped[3:], level=2)
                i += 1
                continue
            elif line_stripped.startswith('### '):
                p = doc.add_heading(line_stripped[4:], level=3)
                i += 1
                continue
            elif line_stripped.startswith('#### '):
                p = doc.add_heading(line_stripped[5:], level=4)
                i += 1
                continue
            elif line_stripped.startswith('##### '):
                p = doc.add_heading(line_stripped[6:], level=5)
                i += 1
                continue
            elif line_stripped.startswith('###### '):
                p = doc.add_heading(line_stripped[7:], level=6)
                i += 1
                continue
            
            # Горизонтальная линия
            if line_stripped.startswith('---') or line_stripped.startswith('***'):
                p = doc.add_paragraph()
                p.add_run('_' * 50)
                i += 1
                continue
            
            # Нумерованный список
            num_match = re.match(r'^(\d+)\.\s+(.+)$', line_stripped)
            if num_match:
                text = num_match.group(2)
                p = doc.add_paragraph(text, style='List Number')
                # Обрабатываем форматирование внутри
                process_inline_formatting(p, text, doc, temp_images)
                i += 1
                continue
            
            # Маркированный список
            if line_stripped.startswith('- ') or line_stripped.startswith('* '):
                text = line_stripped[2:]
                p = doc.add_paragraph(text, style='List Bullet')
                process_inline_formatting(p, text, doc, temp_images)
                i += 1
                continue
            
            # Обычный параграф
            p = doc.add_paragraph()
            process_inline_formatting(p, line_stripped, doc, temp_images)
            i += 1
        
        # Сохраняем документ
        doc.save(output_file)
        
        # Удаляем временные изображения
        for img_path in temp_images:
            try:
                if os.path.exists(img_path):
                    os.unlink(img_path)
            except:
                pass
        
        print(f"[OK] Файл {output_file} успешно создан!")
        print(f"[INFO] Формулы LaTeX конвертированы в изображения ({len(temp_images)} формул).")
        return True
        
    except ImportError as e:
        print(f"Ошибка: Не установлена библиотека: {e}")
        print("\nУстановите необходимые библиотеки:")
        print("  pip install python-docx matplotlib")
        print("\nИли для лучшего качества формул:")
        print("  pip install pypandoc")
        print("  # И установите pandoc: https://pandoc.org/installing.html")
        return False
    except Exception as e:
        print(f"Ошибка при конвертации: {e}")
        import traceback
        traceback.print_exc()
        # Удаляем временные файлы при ошибке
        for img_path in temp_images:
            try:
                if os.path.exists(img_path):
                    os.unlink(img_path)
            except:
                pass
        return False

def process_inline_formatting(paragraph, text, doc, temp_images):
    """Обрабатывает inline форматирование (жирный, курсив, формулы)"""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Разбиваем текст на части с учетом форматирования
    parts = []
    i = 0
    
    while i < len(text):
        # Inline формула LaTeX (одиночный $)
        if text[i] == '$' and (i + 1 >= len(text) or text[i+1] != '$'):
            end = text.find('$', i + 1)
            if end != -1:
                formula = text[i+1:end]
                parts.append(('formula_inline', formula))
                i = end + 1
                continue
        
        # Блочная формула в тексте (двойной $) - не должно быть в inline, но на всякий случай
        if text[i:i+2] == '$$':
            end = text.find('$$', i + 2)
            if end != -1:
                formula = text[i+2:end]
                parts.append(('formula_block', formula))
                i = end + 2
                continue
        
        # Жирный текст
        if text[i:i+2] == '**':
            end = text.find('**', i + 2)
            if end != -1:
                bold_text = text[i+2:end]
                parts.append(('bold', bold_text))
                i = end + 2
                continue
        
        # Курсив
        if text[i] == '*' and (i == 0 or text[i-1] != '*') and (i+1 >= len(text) or text[i+1] != '*'):
            end = text.find('*', i + 1)
            if end != -1 and (end+1 >= len(text) or text[end+1] != '*'):
                italic_text = text[i+1:end]
                parts.append(('italic', italic_text))
                i = end + 1
                continue
        
        # Код
        if text[i] == '`':
            end = text.find('`', i + 1)
            if end != -1:
                code_text = text[i+1:end]
                parts.append(('code', code_text))
                i = end + 1
                continue
        
        # Обычный текст
        start = i
        while i < len(text):
            if text[i] in ['$', '*', '`']:
                break
            i += 1
        if i > start:
            parts.append(('normal', text[start:i]))
    
    # Добавляем части в параграф
    for part_type, part_text in parts:
        if part_type == 'formula_inline':
            # Inline формула - создаем маленькое изображение
            img_path = latex_to_image(part_text, is_block=False)
            if img_path:
                temp_images.append(img_path)
                run = paragraph.add_run()
                run.add_picture(img_path, width=Inches(2))
            else:
                # Если не удалось, добавляем как текст
                run = paragraph.add_run(part_text)
                run.font.name = 'Cambria Math'
                run.font.size = Pt(11)
        elif part_type == 'formula_block':
            # Блочная формула в тексте (не должно быть, но обработаем)
            img_path = latex_to_image(part_text, is_block=True)
            if img_path:
                temp_images.append(img_path)
                run = paragraph.add_run()
                run.add_picture(img_path, width=Inches(4))
            else:
                run = paragraph.add_run(part_text)
                run.font.name = 'Cambria Math'
                run.font.size = Pt(12)
        else:
            run = paragraph.add_run(part_text)
            if part_type == 'bold':
                run.bold = True
            elif part_type == 'italic':
                run.italic = True
            elif part_type == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            # normal - без изменений

if __name__ == '__main__':
    success = convert_markdown_to_docx()
    sys.exit(0 if success else 1)
